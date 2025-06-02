import os
import time
import datetime
import webbrowser
import io
import base64
import streamlit as st
import requests  # For web search and API calls
import google.generativeai as genai
import difflib  # For document comparison
from config import GEMINI_API_KEY, FINHUB_API_KEY, WEATHER_API_KEY, NEWS_API_KEY
from docx import Document
import fitz  # PyMuPDF for PDF files
import pandas as pd
import speech_recognition as sr
from pydub import AudioSegment
from gtts import gTTS  # For text-to-speech conversion
import streamlit.components.v1 as components
from fpdf import FPDF  # For PDF generation
from youtube_transcript_api import YouTubeTranscriptApi  # For YouTube transcripts
from urllib.parse import urlparse, parse_qs
from PIL import Image  # For image processing
import pytesseract  # For OCR
from textblob import TextBlob  # For NLP analysis
from streamlit_quill import st_quill  # Rich text editor component

# --- Configure Gemini API ---
genai.configure(api_key=GEMINI_API_KEY)

# --- DuckDuckGo Smart Web Search Helper ---
def perform_duckduckgo_search(query):
    """Fetch quick results from DuckDuckGo Instant Answer API."""
    try:
        url = f"https://api.duckduckgo.com/?q={query}&format=json&no_html=1&skip_disambig=1"
        response = requests.get(url)
        if response.status_code == 200:
            data = response.json()
            if data.get("AbstractText"):
                return data["AbstractText"]
            else:
                return "No instant answer available."
        else:
            return "Error fetching search results."
    except Exception as e:
        return f"Search error: {str(e)}"

# --- YouTube Transcript and Summarization Helpers ---
def extract_youtube_id(url):
    """Extracts the YouTube video ID from a URL."""
    parsed_url = urlparse(url)
    if parsed_url.hostname in ['youtu.be']:
        return parsed_url.path[1:]
    elif parsed_url.hostname in ['www.youtube.com', 'youtube.com']:
        if parsed_url.path == '/watch':
            query = parse_qs(parsed_url.query)
            return query.get('v', [None])[0]
        elif parsed_url.path.startswith('/embed/'):
            return parsed_url.path.split('/')[2]
        elif parsed_url.path.startswith('/v/'):
            return parsed_url.path.split('/')[2]
    return None

def get_youtube_transcript(video_url):
    """Fetches the transcript for a given YouTube video URL."""
    video_id = extract_youtube_id(video_url)
    if not video_id:
        st.error("Invalid YouTube URL")
        return None
    try:
        transcript_data = YouTubeTranscriptApi.get_transcript(video_id)
        transcript = " ".join([segment['text'] for segment in transcript_data])
        return transcript
    except Exception as e:
        st.error(f"Error fetching transcript: {e}")
        return None

def summarize_transcript(transcript):
    """Summarizes a YouTube transcript using a chat-based prompt."""
    prompt = f"Summarize the following YouTube video transcript into key points:\n\n{transcript}\n\nSummary:"
    return chat(prompt)

# --- Image Text Extraction Helper ---
def extract_text_from_image(image_file):
    """Extracts text from an image file using pytesseract OCR."""
    try:
        image = Image.open(image_file)
        text = pytesseract.image_to_string(image)
        return text
    except Exception as e:
        st.error(f"Error extracting text: {e}")
        return None

# --- Advanced NLP Analysis Helper ---
def advanced_nlp_analysis(text):
    """
    Performs sentiment analysis and keyword extraction using TextBlob,
    and generates a concise summary using the chat-based approach.
    """
    blob = TextBlob(text)
    sentiment = blob.sentiment  # Contains polarity and subjectivity
    keywords = list(blob.noun_phrases)
    summary = chat(f"Please summarize the following text concisely:\n\n{text}")
    return sentiment, keywords, summary

# --- New Helper Function for Entity Extraction ---
def extract_entity(query, instruction):
    """Uses Gemini to extract folder/file names from natural language queries."""
    try:
        prompt = f"{instruction}\nQuery: {query}\nExtracted name:"
        model = genai.GenerativeModel("gemini-2.0-flash")
        response = model.generate_content(prompt)
        cleaned_response = response.text.strip().strip('"').strip("'").split('\n')[0]
        return cleaned_response
    except Exception as e:
        st.error(f"Extraction error: {str(e)}")
        return None

# --- Safe Rerun Helper ---
def safe_rerun():
    if hasattr(st, "experimental_rerun"):
        st.experimental_rerun()
    elif hasattr(st, "rerun"):
        st.rerun()

# --- Helper Functions for File Reading ---
def read_docx_file(file_path_or_buffer):
    try:
        doc = Document(file_path_or_buffer)
        content = [para.text for para in doc.paragraphs]
        return '\n'.join(content)
    except Exception as e:
        st.error(f"Error reading DOCX file: {e}")
        return None

def read_pdf_file(file_path_or_buffer):
    try:
        if hasattr(file_path_or_buffer, "read"):
            file_bytes = file_path_or_buffer.read()
            doc = fitz.open(stream=file_bytes, filetype="pdf")
        else:
            doc = fitz.open(file_path_or_buffer)
        content = []
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            content.append(page.get_text())
        return '\n'.join(content)
    except Exception as e:
        st.error(f"Error reading PDF file: {e}")
        return None

def read_txt_file(file_path_or_buffer):
    try:
        if hasattr(file_path_or_buffer, "read"):
            return file_path_or_buffer.read().decode("utf-8")
        else:
            with open(file_path_or_buffer, "r", encoding="utf-8") as file:
                return file.read()
    except Exception as e:
        st.error(f"Error reading TXT file: {e}")
        return None

def read_csv_file(file_path_or_buffer):
    try:
        if hasattr(file_path_or_buffer, "read"):
            df = pd.read_csv(file_path_or_buffer)
        else:
            df = pd.read_csv(file_path_or_buffer)
        return df.to_string()
    except Exception as e:
        st.error(f"Error reading CSV file: {e}")
        return None

# Helper for uploaded files (file-like objects)
def read_uploaded_file(uploaded_file):
    file_extension = os.path.splitext(uploaded_file.name)[1].lower()
    if file_extension == '.docx':
        return read_docx_file(uploaded_file)
    elif file_extension == '.pdf':
        return read_pdf_file(uploaded_file)
    elif file_extension == '.txt':
        return read_txt_file(uploaded_file)
    elif file_extension == '.csv':
        return read_csv_file(uploaded_file)
    else:
        st.error(f"Unsupported file type: {file_extension}")
        return None

def read_file(file_path):
    file_extension = os.path.splitext(file_path)[1].lower()
    if file_extension == '.docx':
        return read_docx_file(file_path)
    elif file_extension == '.pdf':
        return read_pdf_file(file_path)
    elif file_extension == '.txt':
        return read_txt_file(file_path)
    elif file_extension == '.csv':
        return read_csv_file(file_path)
    else:
        st.error(f"Unsupported file type: {file_extension}")
        return None

def file_exists(file_path):
    return os.path.exists(file_path)

def get_time():
    now = datetime.datetime.now()
    return f"The current time is {now.strftime('%H:%M:%S')}"

def find_file(file_name, search_root="C:\\"):
    for root, dirs, files in os.walk(search_root):
        if file_name in files:
            return os.path.join(root, file_name)
    return None

def find_directory(dir_name, search_root="C:\\"):
    for root, dirs, files in os.walk(search_root):
        if dir_name in dirs:
            return os.path.join(root, dir_name)
    return None

def list_files_in_directory(directory_path):
    try:
        return os.listdir(directory_path)
    except Exception as e:
        return f"Error listing files in directory: {str(e)}"

def list_scholarship_files(directory_path):
    files = list_files_in_directory(directory_path)
    if not isinstance(files, list):
        return files
    result = []
    for item in files:
        result.append(f"* *{item}*")
    return "\n".join(result)

def transcribe_audio(audio_file):
    recognizer = sr.Recognizer()
    if not audio_file.name.lower().endswith('.wav'):
        try:
            audio = AudioSegment.from_file(audio_file)
            audio_file_wav = io.BytesIO()
            audio.export(audio_file_wav, format="wav")
            audio_file_wav.seek(0)
            audio_file = audio_file_wav
        except Exception as e:
            st.error(f"Error converting audio file: {e}")
            return None
    try:
        with sr.AudioFile(audio_file) as source:
            audio_data = recognizer.record(source)
            return recognizer.recognize_google(audio_data)
    except Exception as e:
        st.error(f"Error transcribing audio: {e}")
        return None

def transcribe_microphone():
    r = sr.Recognizer()
    try:
        with sr.Microphone() as source:
            with st.spinner("Listening..."):
                audio = r.listen(source, timeout=5)
        try:
            text = r.recognize_google(audio)
            return text
        except sr.UnknownValueError:
            st.error("Could not understand audio")
            return None
        except sr.RequestError as e:
            st.error(f"Error: {e}")
            return None
    except OSError:
        st.error("No microphone detected. Please ensure a microphone is connected.")
        return None

def play_audio_with_delay(audio_bytes, delay=10000):
    """
    Converts audio bytes to base64 and creates an HTML audio element that auto-plays after a delay.
    Also listens for key events: "p" to play, "s" to stop.
    Returns the HTML string so it can be stored and rendered persistently.
    """
    b64_audio = base64.b64encode(audio_bytes.read()).decode()
    audio_html = f"""
    <audio id="tts-audio" controls src="data:audio/mp3;base64,{b64_audio}"></audio>
    <script>
    // Auto-play after delay
    setTimeout(function() {{
       var audio = document.getElementById("tts-audio");
       if(audio) {{
          audio.play();
       }}
    }}, {delay});

    // Listen for key events: p to play, s to pause
    document.addEventListener('keydown', function(e) {{
      var audio = document.getElementById("tts-audio");
      if(!audio) return;
      if(e.key === 'p' || e.key === 'P') {{
         audio.play();
      }} else if(e.key === 's' || e.key === 'S') {{
         audio.pause();
      }}
    }});
    </script>
    """
    return audio_html

# --- Document Comparison Helper ---
def compare_documents(doc1, doc2):
    """Generates a unified diff between two document texts."""
    lines1 = doc1.splitlines(keepends=True)
    lines2 = doc2.splitlines(keepends=True)
    diff = difflib.unified_diff(lines1, lines2, fromfile='Document1', tofile='Document2', lineterm='')
    return ''.join(diff)

# --- Conversion Helper Functions ---
def text_to_pdf(text):
    """Converts plain text into a PDF file using FPDF."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for line in text.split('\n'):
        pdf.multi_cell(0, 10, line)
    pdf_bytes = io.BytesIO()
    pdf.output(pdf_bytes)
    pdf_bytes.seek(0)
    return pdf_bytes

def convert_pdf_to_docx(pdf_file):
    """Converts a PDF file to a DOCX file by extracting text and writing paragraphs."""
    text = read_pdf_file(pdf_file)
    if text is None:
        return None
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes

def convert_docx_to_pdf(docx_file):
    """Converts a DOCX file to PDF by extracting text and writing it into a PDF."""
    text = read_docx_file(docx_file)
    if text is None:
        return None
    return text_to_pdf(text)

def convert_text_to_pdf(file):
    """Converts an HTML/Markdown (or TXT) file to PDF by reading its text and converting it."""
    file_extension = os.path.splitext(file.name)[1].lower()
    if file_extension in ['.html', '.htm', '.md', '.markdown', '.txt']:
        try:
            text = file.read().decode("utf-8")
        except Exception as e:
            st.error(f"Error reading file: {e}")
            return None
        return text_to_pdf(text)
    else:
        st.error("Unsupported file type for HTML/Markdown conversion")
        return None

def convert_csv_to_excel(csv_file):
    """Converts a CSV file to Excel format using pandas."""
    try:
        df = pd.read_csv(csv_file)
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False)
        output.seek(0)
        return output
    except Exception as e:
        st.error(f"Error converting CSV to Excel: {e}")
        return None

def convert_excel_to_csv(excel_file):
    """Converts an Excel file to CSV format using pandas."""
    try:
        df = pd.read_excel(excel_file)
        csv_data = df.to_csv(index=False)
        csv_bytes = io.BytesIO(csv_data.encode('utf-8'))
        return csv_bytes
    except Exception as e:
        st.error(f"Error converting Excel to CSV: {e}")
        return None

# --- Gemini Chat Function ---
def chat(query):
    try:
        model = genai.GenerativeModel("gemini-2.0-flash")
        response = model.generate_content(query)
        return response.text.strip()
    except Exception as e:
        return f"Error generating response: {str(e)}"

# --- AI-Driven Insights Helper ---
def get_contextual_insights():
    """
    Analyzes the conversation history stored in st.session_state.messages and returns
    contextual recommendations based on past interactions.
    """
    if not st.session_state.messages or len(st.session_state.messages) < 2:
        return ""  # Not enough context to generate insights.
    
    # Compile conversation history into a single string
    conversation_history = "\n".join(
        [f"{speaker}: {message}" for speaker, message in st.session_state.messages]
    )
    
    prompt = (
        "Based on the following conversation, provide some contextual recommendations or "
        "suggestions for further actions that might be helpful for the user. Keep the advice clear "
        "and concise.\n\n"
        f"{conversation_history}\n\nRecommendations:"
    )
    insights = chat(prompt)
    return insights

# --- New Helper Functions for Weather, News, and Stock APIs ---
def get_weather(location):
    """Fetches weather information for a given location using OpenWeatherMap API."""
    try:
        url = f"http://api.openweathermap.org/data/2.5/weather?q={location}&appid={WEATHER_API_KEY}&units=metric"
        response = requests.get(url)
        if response.status_code == 200:
            data = response.json()
            description = data["weather"][0]["description"]
            temp = data["main"]["temp"]
            humidity = data["main"]["humidity"]
            return f"Weather in {location}:\nDescription: {description.capitalize()}\nTemperature: {temp}°C\nHumidity: {humidity}%"
        else:
            return f"Error fetching weather data: {response.status_code}"
    except Exception as e:
        return f"Error: {str(e)}"

def get_stock(stock_symbol):
    """Fetches stock data for a given symbol using Finnhub API."""
    try:
        url = f"https://finnhub.io/api/v1/quote?symbol={stock_symbol}&token={FINHUB_API_KEY}"
        response = requests.get(url)
        if response.status_code == 200:
            data = response.json()
            current = data.get("c")
            high = data.get("h")
            low = data.get("l")
            open_price = data.get("o")
            previous_close = data.get("pc")
            return (
                f"Stock {stock_symbol}:\n"
                f"Current Price: {current}\n"
                f"High: {high}\n"
                f"Low: {low}\n"
                f"Open: {open_price}\n"
                f"Previous Close: {previous_close}"
            )
        else:
            return f"Error fetching stock data: {response.status_code}"
    except Exception as e:
        return f"Error: {str(e)}"

def get_news():
    """Fetches top news headlines using NewsAPI."""
    try:
        url = f"https://newsapi.org/v2/top-headlines?country=us&apiKey={NEWS_API_KEY}"
        response = requests.get(url)
        if response.status_code == 200:
            data = response.json()
            articles = data.get("articles", [])
            news_str = ""
            for article in articles[:5]:
                news_str += f"Title: {article.get('title')}\nSource: {article.get('source', {}).get('name')}\n\n"
            return news_str if news_str else "No news available."
        else:
            return f"Error fetching news: {response.status_code}"
    except Exception as e:
        return f"Error: {str(e)}"

# --- Modified Process Query Function ---
def process_query(query):
    st.session_state.messages.append(("You", query))
    with st.spinner("Astra is thinking..."):
        time.sleep(0.5)
        lower_query = query.lower()
        response = ""
        
        # Handle core commands
        if "open youtube" in lower_query:
            webbrowser.open("https://www.youtube.com")
            response = "Opening YouTube..."
        elif "open google" in lower_query:
            webbrowser.open("https://www.google.com")
            response = "Opening Google..."
        elif "open wikipedia" in lower_query:
            webbrowser.open("https://www.wikipedia.org")
            response = "Opening Wikipedia..."
        elif "time" in lower_query:
            response = get_time()
            
        # Handle list files requests
        elif "list" in lower_query and "files" in lower_query:
            folder_instruction = """Extract the folder/file name from queries about listing files.
Return ONLY the name without additional text. Examples:
- Query: 'Show files in Documents' -> 'Documents'"""
            folder_name = extract_entity(query, folder_instruction)
            if not folder_name:
                response = "Could not determine folder name from your query."
            else:
                if os.path.isabs(folder_name):
                    target_dir = folder_name
                else:
                    target_dir = find_directory(folder_name)
                if not target_dir:
                    response = f"Folder '{folder_name}' not found!"
                else:
                    files = list_files_in_directory(target_dir)
                    if isinstance(files, list):
                        formatted = "\n".join([f"* *{f}*" for f in files])
                        response = f"Files in *{target_dir}*:\n{formatted}"
                    else:
                        response = files
        
        # Handle data extraction requests
        elif "extract" in lower_query and "data" in lower_query:
            file_instruction = """Extract the file/folder name from data extraction queries.
Return ONLY the name without additional text. Examples:
- Query: 'Get data from report.pdf' -> 'report.pdf'"""
            file_name = extract_entity(query, file_instruction)
            if not file_name:
                response = "Could not determine file name from your query."
            else:
                if os.path.isabs(file_name):
                    target_file = file_name
                else:
                    target_file = find_file(file_name)
                if not target_file:
                    response = f"File '{file_name}' not found!"
                else:
                    content = read_file(target_file)
                    response = f"{target_file} contents:\n\n{content}" if content else "Error reading file"
        
        # General Gemini response
        else:
            response = chat(query)
        
        # Append the main response
        st.session_state.messages.append(("Astra", response))
        
        # --- Append AI-Driven Insights ---
        insights = get_contextual_insights()
        if insights:
            response_with_insights = f"{response}\n\n**Additional Recommendations:**\n{insights}"
            # Update the message with the insights
            st.session_state.messages[-1] = ("Astra", response_with_insights)
            response = response_with_insights
        
        # TTS Handling (if enabled)
        if st.session_state.get("read_aloud", False):
            try:
                tts = gTTS(response)
                audio_bytes = io.BytesIO()
                tts.write_to_fp(audio_bytes)
                audio_bytes.seek(0)
                audio_html = play_audio_with_delay(audio_bytes, delay=2000)
                st.session_state.tts_audio = audio_html
            except Exception as e:
                st.error(f"Error in text-to-speech conversion: {e}")
        
        st.session_state.user_input = ""
        safe_rerun()

# --- Main Function ---
def main():
    if 'messages' not in st.session_state:
        st.session_state.messages = []
    if 'user_input' not in st.session_state:
        st.session_state.user_input = ""
    if 'tts_audio' not in st.session_state:
        st.session_state.tts_audio = ""

    st.markdown(
        """
        <style>
        .chat-container {
            height: 400px;
            overflow-y: auto;
            padding: 10px;
            border: 1px solid #ddd;
            background-color: #f9f9f9;
            border-radius: 5px;
        }
        .message {
            margin-bottom: 10px;
            padding: 8px 12px;
            border-radius: 10px;
            max-width: 80%;
        }
        .message-you {
            background-color: #d1e7dd;
            align-self: flex-end;
        }
        .message-astra {
            background-color: #cff4fc;
            align-self: flex-start;
        }
        .message p {
            margin: 0;
        }
        .stButton>button {
            width: 100%;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    with st.sidebar:
        st.header("Astra S-1 Assistant")
        st.info(
            """
            *Instructions:*
            - Use the radio button below to select a mode.
            - **Chat Mode:** Ask questions and control file actions.
            - **Document Comparison Mode:** Upload two documents to compare.
            - **File Conversion Mode:** Convert files between formats.
            - **Web Search Mode:** Get quick results from DuckDuckGo.
            - **YouTube Summarization Mode:** Summarize key points from a YouTube video.
            - **Image Text Extraction Mode:** Extract text from uploaded image files.
            - **Advanced NLP Analysis Mode:** Perform sentiment analysis, keyword extraction, and summarization.
            - **Rich Text Editor Mode:** Enter rich text with formatting.
            - **Weather Mode:** Get current weather information.
            - **News Mode:** Fetch top news headlines.
            - **Stock Mode:** Retrieve stock data using Finnhub.
            """
        )
        if st.button("Clear Conversation"):
            st.session_state.messages = []
            safe_rerun()

        st.checkbox("Read bot responses aloud", key="read_aloud", value=False)
        uploaded_audio = st.file_uploader("Upload Audio File", type=["wav", "mp3", "ogg", "m4a"])
        if uploaded_audio:
            transcribed_text = transcribe_audio(uploaded_audio)
            if transcribed_text:
                st.session_state.user_input = transcribed_text
                safe_rerun()

        mode = st.radio("Select Mode", [
            "Chat", 
            "Document Comparison", 
            "File Conversion", 
            "Web Search", 
            "YouTube Summarization", 
            "Image Text Extraction",
            "Advanced NLP Analysis",
            "Rich Text Editor",
            "Weather",
            "News",
            "Stock"
        ])
        
        if mode == "Document Comparison":
            st.markdown("### Document Comparison")
            comp_file1 = st.file_uploader("Upload Document 1", key="comp_file1", type=["docx", "pdf", "txt", "csv"])
            comp_file2 = st.file_uploader("Upload Document 2", key="comp_file2", type=["docx", "pdf", "txt", "csv"])
            if comp_file1 and comp_file2:
                if st.button("Compare Documents"):
                    doc1_content = read_uploaded_file(comp_file1)
                    doc2_content = read_uploaded_file(comp_file2)
                    if doc1_content is None or doc2_content is None:
                        st.error("Error reading one of the documents.")
                    else:
                        diff_result = compare_documents(doc1_content, doc2_content)
                        st.text_area("Comparison Result (Unified Diff)", diff_result, height=300)
                        
        if mode == "File Conversion":
            st.markdown("### File Conversion")
            conversion_option = st.radio("Select Conversion Type", 
                                          ["PDF to DOCX", "DOCX to PDF", "HTML/Markdown to PDF", "CSV to Excel", "Excel to CSV"])
            if conversion_option in ["PDF to DOCX", "DOCX to PDF", "HTML/Markdown to PDF", "CSV to Excel", "Excel to CSV"]:
                upload_label = "Upload file for conversion"
                conv_file = st.file_uploader(upload_label, key="conv_file", 
                                             type={
                                                 "PDF to DOCX": ["pdf"],
                                                 "DOCX to PDF": ["docx"],
                                                 "HTML/Markdown to PDF": ["html", "htm", "md", "markdown", "txt"],
                                                 "CSV to Excel": ["csv"],
                                                 "Excel to CSV": ["xlsx", "xls"]
                                             }[conversion_option])
                if conv_file:
                    if st.button("Convert File"):
                        converted_bytes = None
                        file_name = conv_file.name
                        if conversion_option == "PDF to DOCX":
                            converted_bytes = convert_pdf_to_docx(conv_file)
                            out_ext = ".docx"
                        elif conversion_option == "DOCX to PDF":
                            converted_bytes = convert_docx_to_pdf(conv_file)
                            out_ext = ".pdf"
                        elif conversion_option == "HTML/Markdown to PDF":
                            converted_bytes = convert_text_to_pdf(conv_file)
                            out_ext = ".pdf"
                        elif conversion_option == "CSV to Excel":
                            converted_bytes = convert_csv_to_excel(conv_file)
                            out_ext = ".xlsx"
                        elif conversion_option == "Excel to CSV":
                            converted_bytes = convert_excel_to_csv(conv_file)
                            out_ext = ".csv"
                        if converted_bytes is None:
                            st.error("Conversion failed.")
                        else:
                            st.download_button("Download Converted File", data=converted_bytes,
                                               file_name=f"converted{out_ext}")
                                               
        if mode == "Web Search":
            st.markdown("### Smart Web Search")
            search_query = st.text_input("Enter search query", key="search_query")
            if st.button("Search"):
                result = perform_duckduckgo_search(search_query)
                st.text_area("Search Results", result, height=300)
                
        if mode == "YouTube Summarization":
            st.markdown("### YouTube Video Summarization")
            youtube_url = st.text_input("Enter YouTube Video URL", key="youtube_url")
            if youtube_url and st.button("Summarize Video"):
                transcript = get_youtube_transcript(youtube_url)
                if transcript:
                    summary = summarize_transcript(transcript)
                    st.text_area("Video Summary", summary, height=300)
                    
        if mode == "Image Text Extraction":
            st.markdown("### Image Text Extraction")
            img_file = st.file_uploader("Upload an Image", key="img_file", type=["png", "jpg", "jpeg", "tiff", "bmp"])
            if img_file and st.button("Extract Text"):
                extracted_text = extract_text_from_image(img_file)
                if extracted_text:
                    st.text_area("Extracted Text", extracted_text, height=300)
                    
        if mode == "Advanced NLP Analysis":
            st.markdown("### Advanced NLP Analysis")
            nlp_text = st.text_area("Enter text for analysis", key="nlp_text", height=200)
            if nlp_text and st.button("Analyze Text"):
                sentiment, keywords, summary = advanced_nlp_analysis(nlp_text)
                st.markdown(f"**Sentiment Analysis:**\n- Polarity: {sentiment.polarity:.3f}\n- Subjectivity: {sentiment.subjectivity:.3f}")
                st.markdown(f"**Extracted Keywords:**\n{', '.join(keywords)}")
                st.text_area("Summary", summary, height=200)
                
        if mode == "Rich Text Editor":
            st.markdown("### Rich Text Editor")
            rich_text = st_quill("Enter rich text below:")
            if rich_text:
                st.markdown("**Your Rich Text Content:**", unsafe_allow_html=True)
                st.markdown(rich_text, unsafe_allow_html=True)
                
        if mode == "Weather":
            st.markdown("### Weather Information")
            location = st.text_input("Enter location (city name)", key="weather_location")
            if location and st.button("Get Weather"):
                weather_info = get_weather(location)
                st.text_area("Weather Details", weather_info, height=150)
                
        if mode == "News":
            st.markdown("### Top News Headlines")
            if st.button("Get News"):
                news_info = get_news()
                st.text_area("News Headlines", news_info, height=300)
                
        if mode == "Stock":
            st.markdown("### Stock Information")
            stock_symbol = st.text_input("Enter stock symbol (e.g., AAPL)", key="stock_symbol")
            if stock_symbol and st.button("Get Stock Data"):
                stock_info = get_stock(stock_symbol)
                st.text_area("Stock Details", stock_info, height=150)

    st.title("Astra S-1: Your Personal AI Assistant")
    
    if mode == "Chat":
        chat_container = st.container()
        with chat_container:
            if st.session_state.messages:
                for speaker, message in st.session_state.messages:
                    if speaker == "Astra":
                        st.markdown(
                            f'<div class="message message-astra"><p><strong>Astra:</strong> {message}</p></div>',
                            unsafe_allow_html=True,
                        )
                    else:
                        st.markdown(
                            f'<div class="message message-you" style="text-align: right;"><p><strong>You:</strong> {message}</p></div>',
                            unsafe_allow_html=True,
                        )
            else:
                st.info("Your conversation will appear here.")
        if st.session_state.tts_audio:
            components.html(st.session_state.tts_audio, height=200)
        col1, col2 = st.columns([5, 1])
        with col1:
            with st.form(key='text_form', clear_on_submit=True):
                user_input = st.text_input(
                    "Type or speak your query",
                    value=st.session_state.user_input,
                    key="user_input_widget",
                    placeholder="Ask me anything..."
                )
                submitted = st.form_submit_button("Send")
        with col2:
            if st.button("🎤", help="Press and speak"):
                transcribed_text = transcribe_microphone()
                if transcribed_text:
                    st.session_state.user_input = transcribed_text
                    safe_rerun()
        if submitted and user_input:
            process_query(user_input)
    
    st.markdown("---")
    st.markdown(
        """
        *Note:* This app runs in a controlled environment. For full file access and advanced features, run locally.
        """
    )

if __name__ == "__main__":
    main()
